"""
Convenience functions for creating word document tables
"""

from typing import List

import docx
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


class DocTable:
    @staticmethod
    def make_cell_bold(cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    def __init__(
        self,
        headers: List[str],
        table_style: str = "Table Grid",
        font: str = "Calibri",
        fontsize: int = 9,
    ) -> None:
        self.doc = docx.Document()
        self.table = self.doc.add_table(rows=1, cols=len(headers))
        self.table.autofit = True
        self.table.allow_autofit = True
        self.table.style = table_style

        self.text_style = self.doc.styles.add_style(
            "My Table Text", WD_STYLE_TYPE.PARAGRAPH
        )
        self.text_style.font.name = font
        self.text_style.font.size = Pt(fontsize)

        self.nrows = 1
        self.ncolumns = len(headers)

        self.header_row = self.table.row_cells(0)

        for idx, header in enumerate(headers):
            p = self.header_row[idx].paragraphs[0]
            p.style = self.text_style
            p.add_run(header).bold = True

    def add_row(self, row_elements: List[str]):
        new_row = self.table.add_row().cells

        for cell in new_row:
            # In case number of row elements doesn't completely fill row
            cell.paragraphs[0].style = self.text_style

        for idx, txt in enumerate(row_elements):
            new_row[idx].paragraphs[0].add_run(txt)

    def rename_rows(self, mapper: dict):
        """
        Locate rows with a specific value in first column and change to specified value
        """
        for original_val, new_val in mapper.items():
            raise NotImplemented

    def save(self, path: str):
        self.doc.save(path)
